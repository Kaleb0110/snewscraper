from bs4 import BeautifulSoup
import urllib.request
from docx import Document
from datetime import datetime
from urllib.error import HTTPError

print('Extracting top news feed xml from The Standard website...')
url = 'https://www.thestandard.com.hk/newsfeed/latest/news.xml'
page = urllib.request.urlopen(url)
soup = BeautifulSoup(page, 'xml')

list_of_news_html = []

print('Extracting content from each top news link...')
for i,e in enumerate(soup.find_all('link')):
    if i == 0 or i == 1:
        continue
    else:
        try:
            news_page = urllib.request.urlopen(e.contents[0])
            news_soup = BeautifulSoup(news_page, 'lxml')
            list_of_news_html.append(news_soup.find_all('p')[:-2])
        except HTTPError:
            print('Skipping over link because of HTTP Error.')

print('Formatting content for document creation...')
news_string = ''
list_of_news = []
for i in list_of_news_html:
    for j,k in enumerate(i):
        news_string += k.contents[0] + '\n'
        if j == (len(i)-1):
            list_of_news.append(news_string)
    news_string = ''

print('Creating document for export...')
news_doc = Document()
news_doc.add_heading('The Standard Top News',0)
today = datetime.now()
date_string = 'Extracted at ' + today.strftime("%d/%m/%Y %H:%M:%S")
news_doc.add_paragraph(date_string)

for i in list_of_news:
    news_doc.add_heading(i.split('\n')[0], 0)
    print(i)
    p = news_doc.add_paragraph(i)

news_doc.save('demo.docx')